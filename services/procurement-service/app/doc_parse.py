"""Extract plain text from an uploaded BOM document (pdf / docx / txt / csv)."""
from __future__ import annotations

import io


def extract_text(filename: str, data: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        return "\n".join((p.extract_text() or "") for p in reader.pages)
    if name.endswith(".docx"):
        import docx
        d = docx.Document(io.BytesIO(data))
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        for t in d.tables:
            for row in t.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    # txt / csv / anything else
    return data.decode("utf-8", errors="ignore")
